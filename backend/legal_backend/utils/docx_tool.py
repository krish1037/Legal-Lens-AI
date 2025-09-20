# import os
# from typing import List
# import docx
# from legal_backend.utils.pdf_tools import extract_from_pdf

# def extract_from_docx(docx_path: str) -> str:
#     """
#     Extract text from a Word document (.docx).
#     Falls back to PDF OCR if no text is found (e.g., scanned DOCX).
#     """
#     if not os.path.exists(docx_path):
#         return f"File not found: {docx_path}"

#     doc = docx.Document(docx_path)
#     paragraphs: List[str] = [p.text for p in doc.paragraphs if p.text.strip()]

#     if paragraphs:
#         return "\n".join(paragraphs).strip()

#     # If no text (could be scanned DOCX with embedded images) → convert to PDF
#     # Requires libreoffice or unoconv installed
#     tmp_pdf = docx_path.replace(".docx", ".pdf")
#     os.system(f'libreoffice --headless --convert-to pdf "{docx_path}" --outdir "{os.path.dirname(docx_path)}"')

#     if os.path.exists(tmp_pdf):
#         return extract_from_pdf(tmp_pdf)

#     return "No text detected."
import os
from typing import List
import docx
from legal_backend.utils.pdf_tools import extract_from_pdf

def extract_from_docx(docx_path: str) -> dict:
    """
    Extract text from a Word document (.docx).
    Falls back to PDF OCR if no text is found (e.g., scanned DOCX).
    Always returns JSON.
    """
    if not os.path.exists(docx_path):
        return {
            "status": "error",
            "source": "docx",
            "text": f"File not found: {docx_path}"
        }

    try:
        doc = docx.Document(docx_path)
        paragraphs: List[str] = [p.text for p in doc.paragraphs if p.text.strip()]

        if paragraphs:
            return {
                "status": "success",
                "source": "docx",
                "text": "\n".join(paragraphs).strip()
            }

        # If no text (e.g., scanned DOCX with embedded images) → convert to PDF
        tmp_pdf = docx_path.replace(".docx", ".pdf")
        os.system(f'libreoffice --headless --convert-to pdf "{docx_path}" --outdir "{os.path.dirname(docx_path)}"')

        if os.path.exists(tmp_pdf):
            pdf_result = extract_from_pdf(tmp_pdf)
            return {
                "status": pdf_result.get("status", "error"),
                "source": "docx",
                "text": pdf_result.get("text", "No text detected after PDF conversion.")
            }

        return {
            "status": "error",
            "source": "docx",
            "text": "No text detected and PDF conversion failed."
        }

    except Exception as e:
        return {
            "status": "error",
            "source": "docx",
            "text": f"Unexpected error: {str(e)}"
        }
